"""
Document ingestion, chunking, and metadata extraction.
"""
import io
from pathlib import Path
from typing import List, Dict, Any, BinaryIO
import pandas as pd
from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

from core.config import Config
from core.logging_utils import get_logger

logger = get_logger(__name__)


class DocumentIngestor:
    """Handles document loading, parsing, and chunking."""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.md', '.csv'}
    
    def __init__(self, config: Config):
        """
        Initialize document ingestor.
        
        Args:
            config: Application configuration
        """
        self.config = config
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.chunk_size,
            chunk_overlap=config.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    def is_supported(self, filename: str) -> bool:
        """
        Check if file type is supported.
        
        Args:
            filename: Name of the file
            
        Returns:
            True if supported, False otherwise
        """
        ext = Path(filename).suffix.lower()
        return ext in self.SUPPORTED_EXTENSIONS
    
    def ingest_file(self, file_path: str = None, file_obj: BinaryIO = None, 
                   filename: str = None) -> List[Document]:
        """
        Ingest a file and return chunked documents with metadata.
        
        Args:
            file_path: Path to file (for file system)
            file_obj: File object (for uploaded files)
            filename: Original filename
            
        Returns:
            List of Document objects with metadata
        """
        if file_path:
            filepath = Path(file_path)
            filename = filename or filepath.name
            ext = filepath.suffix.lower()
            
            with open(filepath, 'rb') as f:
                content = f.read()
        elif file_obj and filename:
            ext = Path(filename).suffix.lower()
            content = file_obj.read()
            if hasattr(file_obj, 'seek'):
                file_obj.seek(0)  # Reset for potential re-reading
        else:
            raise ValueError("Either file_path or (file_obj and filename) must be provided")
        
        if not self.is_supported(filename):
            raise ValueError(f"Unsupported file type: {ext}")
        
        logger.info(f"Ingesting file: {filename}")
        
        # Parse based on file type
        if ext == '.pdf':
            documents = self._parse_pdf(content, filename)
        elif ext == '.docx':
            documents = self._parse_docx(content, filename)
        elif ext in ['.txt', '.md']:
            documents = self._parse_text(content, filename)
        elif ext == '.csv':
            documents = self._parse_csv(content, filename)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
        
        logger.info(f"Created {len(documents)} chunks from {filename}")
        return documents
    
    def _parse_pdf(self, content: bytes, filename: str) -> List[Document]:
        """Parse PDF and extract text with page numbers."""
        documents = []
        
        try:
            pdf_reader = PdfReader(io.BytesIO(content))
            
            for page_num, page in enumerate(pdf_reader.pages, start=1):
                text = page.extract_text()
                if text.strip():
                    # Create a document for each page first
                    doc = Document(
                        page_content=text,
                        metadata={
                            'filename': filename,
                            'filetype': 'pdf',
                            'page': page_num
                        }
                    )
                    
                    # Split into chunks while preserving page metadata
                    chunks = self.text_splitter.split_documents([doc])
                    
                    # Ensure all chunks have the page number
                    for chunk in chunks:
                        chunk.metadata['page'] = page_num
                    
                    documents.extend(chunks)
        
        except Exception as e:
            logger.error(f"Error parsing PDF {filename}: {e}")
            raise
        
        return documents
    
    def _parse_docx(self, content: bytes, filename: str) -> List[Document]:
        """Parse DOCX file."""
        try:
            doc = DocxDocument(io.BytesIO(content))
            text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
            
            base_doc = Document(
                page_content=text,
                metadata={
                    'filename': filename,
                    'filetype': 'docx'
                }
            )
            
            documents = self.text_splitter.split_documents([base_doc])
            
            # Ensure all chunks have metadata
            for doc in documents:
                doc.metadata.update({
                    'filename': filename,
                    'filetype': 'docx'
                })
            
            return documents
        
        except Exception as e:
            logger.error(f"Error parsing DOCX {filename}: {e}")
            raise
    
    def _parse_text(self, content: bytes, filename: str) -> List[Document]:
        """Parse plain text or markdown file."""
        try:
            text = content.decode('utf-8', errors='ignore')
            
            base_doc = Document(
                page_content=text,
                metadata={
                    'filename': filename,
                    'filetype': Path(filename).suffix[1:]  # Remove the dot
                }
            )
            
            documents = self.text_splitter.split_documents([base_doc])
            
            # Ensure all chunks have metadata
            for doc in documents:
                doc.metadata.update({
                    'filename': filename,
                    'filetype': Path(filename).suffix[1:]
                })
            
            return documents
        
        except Exception as e:
            logger.error(f"Error parsing text file {filename}: {e}")
            raise
    
    def _parse_csv(self, content: bytes, filename: str) -> List[Document]:
        """Parse CSV file and convert to text summaries."""
        try:
            df = pd.read_csv(io.BytesIO(content))
            
            # Create a text representation
            # Header information
            text_parts = [
                f"CSV File: {filename}",
                f"Columns: {', '.join(df.columns.tolist())}",
                f"Rows: {len(df)}",
                "\n--- Data ---\n"
            ]
            
            # Add row summaries (limit to prevent huge texts)
            max_rows = min(len(df), 1000)  # Limit to 1000 rows
            for idx, row in df.head(max_rows).iterrows():
                row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
                text_parts.append(f"Row {idx + 1}: {row_text}")
            
            if len(df) > max_rows:
                text_parts.append(f"\n... and {len(df) - max_rows} more rows")
            
            text = '\n'.join(text_parts)
            
            base_doc = Document(
                page_content=text,
                metadata={
                    'filename': filename,
                    'filetype': 'csv',
                    'rows': len(df),
                    'columns': len(df.columns)
                }
            )
            
            documents = self.text_splitter.split_documents([base_doc])
            
            # Ensure all chunks have metadata
            for doc in documents:
                doc.metadata.update({
                    'filename': filename,
                    'filetype': 'csv'
                })
            
            return documents
        
        except Exception as e:
            logger.error(f"Error parsing CSV {filename}: {e}")
            raise
    
    def ingest_multiple(self, file_paths: List[str] = None,
                       file_objects: List[tuple] = None) -> List[Document]:
        """
        Ingest multiple files.
        
        Args:
            file_paths: List of file paths
            file_objects: List of (file_obj, filename) tuples
            
        Returns:
            Combined list of documents
        """
        all_documents = []
        
        if file_paths:
            for path in file_paths:
                try:
                    docs = self.ingest_file(file_path=path)
                    all_documents.extend(docs)
                except Exception as e:
                    logger.error(f"Failed to ingest {path}: {e}")
        
        if file_objects:
            for file_obj, filename in file_objects:
                try:
                    docs = self.ingest_file(file_obj=file_obj, filename=filename)
                    all_documents.extend(docs)
                except Exception as e:
                    logger.error(f"Failed to ingest {filename}: {e}")
        
        return all_documents
